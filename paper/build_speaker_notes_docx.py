"""Generate Prompt_Pattern_Mining_Speaker_Notes.docx in Downloads."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path.home() / "Downloads" / "Prompt_Pattern_Mining_Speaker_Notes.docx"

NOTES = [
    {
        "n": 1, "title": "Title slide",
        "screen": "Paper title, three author names, TAMUSA Department of Computer Science.",
        "say": (
            "Good afternoon everyone. I'm Visesh Bentula, and along with my "
            "teammates Teja Reddy Mandadi and Umapathi Konduri, I'll be "
            "presenting our project: Prompt Pattern Mining in Vibe Coding "
            "- A Statistical Audit of 6,413 Developer Conversations. The work "
            "was done as part of the Graduate Seminar course at Texas A&M "
            "University-San Antonio.\n\n"
            "The basic question we set out to answer is simple: when a "
            "developer sits down to write code with ChatGPT, which prompt "
            "features actually predict whether the conversation produces "
            "useful code? We answered that on a public corpus of 6,413 real "
            "ChatGPT conversations, and we replicated the result on three "
            "other top-tier LLMs. I'll walk you through what we found in the "
            "next twenty minutes."
        ),
        "qa": [(
            "Why call it vibe coding?",
            "It's the slang the industry is now using for the iterative, "
            "prompt-driven way developers write code with conversational "
            "LLMs. We use it because it captures a real shift - the artifact "
            "under study is no longer the code itself, it's the prompt "
            "sequence that produced it."
        )],
    },
    {
        "n": 2, "title": "Problem Statement & Motivation",
        "screen": "Title only - likely with a visual or bullets framing the problem.",
        "say": (
            "Conversational LLMs have moved into the day-to-day workflow of "
            "software development. If you look at modern pull request threads, "
            "GitHub issues, or developer forums, you'll see long ChatGPT "
            "exchanges where the developer iterates on a prompt until the "
            "code looks right.\n\n"
            "This raises a real question for the field: the wording of the "
            "prompt drives the quality of the answer, but we don't yet have "
            "large-scale evidence on which structural prompt features "
            "actually work in real developer workflows.\n\n"
            "A lot of prompt engineering tutorials exist, but most of them "
            "are based on either gut feel or experimenter-written prompts in "
            "artificial benchmarks. Nobody had asked: on thousands of real "
            "developer-ChatGPT conversations that already happened in the "
            "wild, which prompt features actually predict whether the "
            "conversation produces useful code?\n\n"
            "That's the gap. The motivation for this work is to close it "
            "quantitatively."
        ),
        "qa": [(
            "What's wrong with the existing prompt engineering studies?",
            "Two things. First, they use experimenter-authored prompts on "
            "synthetic benchmarks, so the results may not transfer to "
            "natural developer dialogue. Second, the descriptive studies of "
            "natural dialogue, like the DevGPT papers, never fit a "
            "predictive model on the success outcome. Our paper does both."
        )],
    },
    {
        "n": 3, "title": "Research Questions",
        "screen": "Title - likely lists RQ1-RQ4 visually.",
        "say": (
            "We organized the work around four research questions.\n\n"
            "RQ1. Which structural prompt features predict successful coding "
            "conversations after controlling for confounders like programming "
            "language, the GitHub source type, and the snapshot date?\n\n"
            "RQ2. How does iterative behavior - meaning back-and-forth turns "
            "- when separated from corrective behavior, affect success?\n\n"
            "RQ3. Do prompt features still carry predictive load when richer "
            "textual baselines, like a bag-of-words classifier, are "
            "available?\n\n"
            "RQ4. Does the engineered prompt effect generalize beyond ChatGPT "
            "to other top-tier LLMs?\n\n"
            "Each maps to a result section in the paper. By the end of this "
            "talk you'll see a yes-or-no answer for each one, backed by a "
            "statistical test."
        ),
        "qa": [(
            "Why did you pre-register these questions?",
            "Pre-registering means we wrote the questions down before "
            "fitting the model. That stops us from cherry-picking the "
            "questions we ended up answering - a known failure mode in "
            "observational studies."
        )],
    },
    {
        "n": 4, "title": "Proposed Idea",
        "screen": "Title - high-level approach.",
        "say": (
            "Our proposal is to do a quantitative audit of real developer "
            "prompts. The recipe has four parts.\n\n"
            "First, we take a public corpus of real conversations - DevGPT "
            "v10, which contains 6,413 ChatGPT exchanges linked to GitHub "
            "artifacts.\n\n"
            "Second, we extract eleven measurable features from each "
            "conversation using regex rules - things like prompt length, "
            "the presence of an output format instruction, the number of "
            "refinement turns, and so on.\n\n"
            "Third, we define a textual success label - basically, a "
            "conversation counts as success if the assistant produced code "
            "and the developer never wrote 'this doesn't work' or similar "
            "in any later turn.\n\n"
            "Fourth, we fit a logistic regression that predicts the success "
            "label from those eleven features, controlling for things like "
            "programming language, source type, and snapshot date. Then we "
            "run a battery of robustness checks on top of that.\n\n"
            "That's the whole idea. The novelty is that we do this on real "
            "conversations, at scale, with proper statistical controls."
        ),
        "qa": [(
            "Why logistic regression and not a deep model?",
            "Three reasons. It's interpretable - every feature gets one "
            "coefficient and we can read off the effect. It's fast - fits "
            "in seconds on six thousand rows. And it gives us confidence "
            "intervals and p-values for free, which a neural net does not. "
            "We do compare against a bag-of-words baseline to make sure "
            "we're not leaving performance on the table."
        )],
    },
    {
        "n": 5, "title": "Research Methodology",
        "screen": "Title - likely a pipeline diagram.",
        "say": (
            "This slide shows our six-stage pipeline, which is also Figure 1 "
            "in the paper.\n\n"
            "We start with the raw DevGPT v10 JSON. The first stage is "
            "cleaning and tokenization - we parse the JSON, deduplicate, "
            "drop corrupted threads, and normalize encoding.\n\n"
            "Stage two is regex feature extraction. We use a hand-curated "
            "regex catalog, committed to the project repo, to detect things "
            "like role instructions, output format requests, and refinement "
            "language.\n\n"
            "Stage three is automatic success labeling - applying our "
            "textual heuristic to each conversation.\n\n"
            "Stage four is the logistic regression fit itself, with L2 "
            "regularization and class-weight balanced.\n\n"
            "Stage five is the suite of robustness checks: five-fold cross "
            "validation, baseline comparison, ablation, pre-registered "
            "interactions, and a manual re-labeling audit.\n\n"
            "Every stage corresponds to a script in the pipeline/ folder of "
            "our repo, and the random seed is fixed at 42 throughout, so "
            "any of these results can be regenerated by anyone with our "
            "codebase."
        ),
        "qa": [(
            "How do you handle the imbalance between success and failure?",
            "55-45 isn't bad, but we use scikit-learn's class_weight=balanced "
            "option, which weights each class inversely to its frequency. "
            "That prevents the model from cheating by always predicting the "
            "majority class."
        )],
    },
    {
        "n": 6, "title": "Dataset & Data Construction",
        "screen": (
            "DevGPT v10 corpus | 6,413 conversations, 54,449 turns | "
            "Sources: PRs, issues, commits, discussions, Hacker News | "
            "Processing steps: JSON parsing, deduplication, regex extraction, "
            "auto-labeling."
        ),
        "say": (
            "Let me give you the numbers. We use DevGPT v10, which is a "
            "publicly released corpus put out by Xiao and colleagues. After "
            "our cleaning pipeline, we end up with 6,413 conversations and "
            "54,449 turns - about eight and a half turns per conversation "
            "on average.\n\n"
            "The conversations are linked to six different source types: "
            "pull requests, GitHub issues, commits, discussions, code-file "
            "embeds, and Hacker News threads. They span ten snapshots "
            "between August 2023 and May 2024.\n\n"
            "The processing pipeline does five things: it parses the raw "
            "JSON, deduplicates conversations, extracts the eleven features "
            "via regex, applies the automatic success label, and writes a "
            "structured parquet dataset that the modeling step reads. The "
            "whole thing runs end-to-end with one command - python main.py "
            "- and produces all the artifacts referenced in the paper.\n\n"
            "Crucially: 3,529 of the 6,413 conversations are labeled "
            "success and 2,884 are labeled failure. That's a 55-45 split - "
            "close to balanced, which keeps the model honest."
        ),
        "qa": [(
            "How did you confirm the dataset is clean?",
            "Two ways. First, we manually re-labeled 50 stratified random "
            "conversations and compared against the auto-labeler - got "
            "Cohen's kappa 0.834, which is almost perfect agreement. Second, "
            "we control for source type, language, and snapshot date in the "
            "regression, so any residual quality variation across sources "
            "gets absorbed by those covariates."
        )],
    },
    {
        "n": 7, "title": "Experiments",
        "screen": "Title - likely an overview of the experiment families.",
        "say": (
            "Our experimental design has four orthogonal components. Each "
            "one attacks the central question from a different angle, so if "
            "all four agree, we know the result is real and not an artifact "
            "of one method.\n\n"
            "Component one is the headline regression - a balanced "
            "L2-regularized logistic regression with five-fold stratified "
            "cross validation, fit on the full eleven features plus "
            "controls.\n\n"
            "Component two is a baseline comparison. We fit three control "
            "classifiers on the same labels: majority class, random, and a "
            "bag-of-words logistic with two thousand unigram and bigram "
            "features.\n\n"
            "Component three is a feature group ablation - we re-fit the "
            "model five times, each time removing one feature group, and "
            "measure the AUC drop.\n\n"
            "Component four is the cross-model replication. We replay 200 "
            "prompts through three other LLMs - Kimi K2, Claude Sonnet 4.6, "
            "and Gemini 3.1 Pro - under both zero-shot and engineered "
            "conditions.\n\n"
            "On top of all that we have a manual audit, a pre-registered "
            "interaction model, and language and time subgroup analyses. "
            "Every experiment is reproducible from a JSON or parquet "
            "artifact in the repo."
        ),
        "qa": [],
    },
    {
        "n": 8, "title": "Results",
        "screen": "ROC-AUC 0.799 | Accuracy 0.726 | F1 0.755 | Baseline comparison chart.",
        "say": (
            "Here's the headline number. Under five-fold stratified cross "
            "validation, our model achieves ROC-AUC of 0.799, with a "
            "standard deviation of only 0.009 across the folds. Accuracy is "
            "0.726, F1 is 0.755.\n\n"
            "For context: 0.5 means a random classifier; 0.8 to 0.9 is "
            "considered strong; we're sitting at the upper end of decent. "
            "For a problem where the labels are noisy textual heuristics "
            "and the inputs are eleven hand-coded features, that's honestly "
            "impressive.\n\n"
            "The baseline comparison chart on this slide tells the most "
            "important methodological story. Majority and random sit at "
            "chance, which is sanity-checking. The bag-of-words classifier "
            "on the first prompt - using two thousand features - reaches "
            "AUC 0.782. So most of the predictive signal lives in the words "
            "of the first prompt, which makes sense.\n\n"
            "Our hand-engineered eleven-feature model adds 1.7 AUC points "
            "on top of that. Those extra points come from features "
            "bag-of-words can't capture cleanly: refinement turn counts, "
            "role instruction flags, prompt growth - properties of the "
            "whole conversation, not just the first prompt."
        ),
        "qa": [(
            "Why is bag-of-words so close to your model?",
            "Because language and topic are very predictive on their own - "
            "short CSS prompts succeed often, long Rust prompts succeed "
            "less often. Bag-of-words captures that for free. Our 1.7-point "
            "margin is the part that specifically comes from prompt "
            "structure, which is what the paper is actually about. And our "
            "features are interpretable - bag-of-words isn't."
        )],
    },
    {
        "n": 9, "title": "Key Results & Findings",
        "screen": (
            "Output format +75% odds | Refinement +19% per turn | Role alone "
            "hurts | Role x Format = 5.28 | Structured > BoW."
        ),
        "say": (
            "This is the slide where the actual findings live. Four "
            "headline results.\n\n"
            "One. Adding an explicit output format instruction in the first "
            "prompt raises the odds of success by 75 percent. The odds "
            "ratio is 1.75, the 95 percent confidence interval is 1.40 to "
            "2.19, and the p-value is below one in a million. It does not "
            "matter whether you ask for JSON, markdown, or 'return only "
            "valid Python code in a fenced block' - what matters is that "
            "you anchor the output format somewhere.\n\n"
            "Two. Each additional refinement turn raises the odds of "
            "success by 19 percent. Refinement here means turns where the "
            "developer asks for a fix, an extension, or a clarification - "
            "not corrective turns where the model gave a wrong answer.\n\n"
            "Three. A bare role prompt - 'you are a Python expert' on its "
            "own - lowers the odds of success by 35 percent. That's "
            "surprising, and I'll explain why on the next slide.\n\n"
            "Four. Combining a role prompt with an output format produces "
            "an interaction odds ratio of 5.28. So role plus format is "
            "roughly five times more likely to succeed than neither.\n\n"
            "The bag-of-words classifier, for comparison, can't see any of "
            "these features explicitly. It just sees word counts."
        ),
        "qa": [(
            "Are these effects practically meaningful, or just statistically significant?",
            "Both. A 75 percent lift on something as simple as appending "
            "'return JSON' to your prompt is large by any practical "
            "standard, and the p-values are not just below 0.05 - they're "
            "below 1e-6, so noise is not driving the result."
        )],
    },
    {
        "n": 10, "title": "Feature Importance / Ablation",
        "screen": "Title - likely the ablation bar chart from the paper.",
        "say": (
            "This is the ablation analysis. The way it works: we re-fit "
            "the model five separate times, each time removing one group "
            "of features, and record how much the AUC drops.\n\n"
            "The five groups are: our eleven prompt-engineering features, "
            "repository language one-hots, source type one-hots, snapshot "
            "date one-hots, and ChatGPT model variant one-hots.\n\n"
            "The bar chart shows the AUC loss when each group is removed.\n\n"
            "When we remove the eleven prompt features, AUC drops by 0.043 "
            "- the largest drop by far.\n\n"
            "When we remove repository language: 0.013.\n"
            "When we remove source type: 0.016.\n"
            "When we remove snapshot date: 0.001.\n"
            "ChatGPT model variant: even smaller.\n\n"
            "The takeaway is direct: the eleven prompt features carry "
            "roughly three times the predictive load of any control group, "
            "and roughly three times the load of all four control groups "
            "combined. The predictive power lives in the prompts, not in "
            "the metadata.\n\n"
            "This is the single strongest argument that the paper's claims "
            "are real and not artifacts of confounding by language, source, "
            "or time."
        ),
        "qa": [(
            "Why is the snapshot date contribution so small?",
            "Because once you control for prompt features, the temporal "
            "variation has almost nowhere to act. The patterns that worked "
            "in August 2023 still work in May 2024 - the effect is not "
            "driven by some snapshot of the data."
        )],
    },
    {
        "n": 11, "title": "Interaction Effects Between Prompt Features",
        "screen": (
            "Major finding: role alone hurts; role + output format compounds. "
            "Role x Output Format OR = 5.28, strongest interaction. Two "
            "example prompts. Insight: role prompting only works with "
            "explicit output structure."
        ),
        "say": (
            "This is the most interesting result in the paper, and I want "
            "to spend a minute on the worked example.\n\n"
            "We pre-registered four two-way interactions before fitting "
            "any model. Three of the four come back significant. The one "
            "that dominates is role times output format, with an "
            "interaction odds ratio of 5.28.\n\n"
            "What that means in practice: the role main effect is negative "
            "- odds ratio 0.65, so a bare 'you are a Python expert' on its "
            "own lowers the odds by 35 percent. The output format main "
            "effect is positive - odds ratio 1.75. But when both are "
            "present in the same prompt, the joint factor is 0.65 times "
            "1.75 times 5.28, which works out to about 6.0 - roughly nine "
            "times higher than the role-alone prompt.\n\n"
            "The slide shows the example. The less effective prompt is "
            "'You are a Python expert. Help me sort a list.' Role present, "
            "format absent. The more effective version is 'You are a "
            "Python expert. Sort the list and return only valid Python "
            "code in a fenced block.' Role and format both present.\n\n"
            "The lesson, which is the take-home for any developer in the "
            "room: role instructions are not free. They need a concrete "
            "output anchor to pay off. A bare role prompt without "
            "specifying what you want is associated with vague exploratory "
            "questions, which fail more often. Add a format anchor and the "
            "joint effect flips strongly positive."
        ),
        "qa": [(
            "How do you know the role instruction isn't just causing the failure?",
            "Because the interaction analysis tells us role flips positive "
            "once combined with format. If role were causally bad, the "
            "interaction would still be negative. What we're seeing is "
            "user intent: people who write a bare 'you are X' tend to be "
            "in exploratory mode, and that's what fails. We say so "
            "explicitly in the Discussion section."
        )],
    },
    {
        "n": 12, "title": "Cross-Model Evaluation",
        "screen": (
            "Three models: Kimi K2, Claude Sonnet 4.6, Gemini 3.1 Pro. "
            "200 prompts x 3 vendors x 2 conditions. Gains +14 to +24 pp."
        ),
        "say": (
            "To make sure our finding is not just a quirk of ChatGPT, we "
            "ran a cross-model replication.\n\n"
            "We took 200 prompts from DevGPT and replayed each one through "
            "three independent LLMs from three different vendors:\n"
            "Kimi K2 from Moonshot AI, served through NVIDIA NIM,\n"
            "Claude Sonnet 4.6 from Anthropic, served through AWS Bedrock,\n"
            "Gemini 3.1 Pro Preview from Google, served through Vertex AI.\n\n"
            "Each prompt went through twice: once as zero-shot, with a "
            "generic 'you are a helpful coding assistant' system prompt, "
            "and once as engineered, where the system prompt explicitly "
            "asks for constraints, an output format, and complexity "
            "analysis.\n\n"
            "Each output was graded on a four-axis rubric - code present, "
            "addresses task, looks correct, follows structure - for a 0-to-4 "
            "score. We then computed the success rate using a threshold of "
            "three or higher.\n\n"
            "Results: every single vendor improved with engineered prompts. "
            "The success-rate lifts were +23 percentage points for Kimi, "
            "+14 for Claude, +24 for Gemini. Cohen's d effect sizes range "
            "from 0.59 (Claude) to 0.80 (Gemini) - moderate to large on "
            "Cohen's scale.\n\n"
            "The takeaway is the most important external validity statement "
            "in the paper: prompt engineering generalizes across LLM "
            "vendors. This isn't a ChatGPT quirk."
        ),
        "qa": [(
            "Why is Claude's lift smaller than the others?",
            "Because Claude's zero-shot baseline is already higher - 3.03 "
            "out of 4 versus around 2.5 for Kimi and Gemini. There's less "
            "headroom for the engineered prompt to improve."
        )],
    },
    {
        "n": 13, "title": "Validation & Reliability",
        "screen": (
            "Stratified random sample, 50 conversations re-labeled. "
            "Cohen's kappa = 0.834, raw 92%, only 4 disagreements. "
            "'Almost perfect' on Landis & Koch."
        ),
        "say": (
            "This slide validates the auto-labeler - the textual heuristic "
            "we use to assign success or failure to each conversation. If "
            "the auto-labeler were noisy, all our results would be "
            "suspect.\n\n"
            "Here's what we did. We drew a stratified random sample of 200 "
            "conversations - stratified means the sample preserves the "
            "success-failure split of the full corpus. We hand-labeled the "
            "first 50 against a written protocol, blind to the auto-label.\n\n"
            "The agreement statistic is Cohen's kappa, which is the "
            "standard measure of inter-rater agreement above chance. Kappa "
            "runs from 0 to 1, and we use the Landis and Koch scale to "
            "interpret it.\n\n"
            "Our kappa is 0.834, raw agreement is 92 percent, with only 4 "
            "disagreements out of 50. Under Landis and Koch, kappa above "
            "0.81 counts as almost perfect agreement - the highest band on "
            "the scale.\n\n"
            "The four disagreements are all conservative misses - the "
            "auto-labeler was stricter than the human, marking borderline "
            "cases as failure when the human said success. That direction "
            "is good news for our findings, because it means any bias in "
            "the labeler works against finding effects, not for them. The "
            "real effects are at least as strong as we report."
        ),
        "qa": [(
            "Why only 50 hand labels?",
            "Time. Hand-labeling 50 conversations takes a few hours; 500 "
            "takes weeks. We flag a 500-sample blinded two-reviewer "
            "protocol as our top future-work item, exactly because that "
            "would let us also report human-against-human kappa separately "
            "from human-against-auto kappa."
        )],
    },
    {
        "n": 14, "title": "Discussion",
        "screen": "Title - likely interpretation bullets.",
        "say": (
            "Let me pull the threads together.\n\n"
            "The fitted model gives a short evidence-based checklist for "
            "any developer working with an LLM.\n\n"
            "First, add an output format instruction. A trailing 'return "
            "JSON' or 'respond in markdown with fenced code' alone raises "
            "the odds of a useful answer by 75 percent.\n\n"
            "Second, expect to refine. Each refinement turn adds 19 "
            "percent to the odds of success - provided the conversation "
            "does not contain explicit corrective feedback. Refinement is "
            "part of the workflow, not a sign of trouble. That's a new "
            "finding, because earlier descriptive studies treated all "
            "back-and-forth as bad.\n\n"
            "Third, do not use a bare role prompt. 'You are a Python "
            "expert' on its own slightly hurts. Combine it with an output "
            "format and the joint odds rise by a factor of about five.\n\n"
            "Fourth, keep the first prompt short and specific. Long "
            "prompts and growing prompts both correlate with worse "
            "outcomes - probably reflecting underspecified intent.\n\n"
            "The reason refinement comes out positive in our model, even "
            "though prior studies said it was negative, is that we "
            "separate refinement from correction cycles. Refinement is "
            "the developer iterating on a working answer; correction "
            "cycles are the developer fixing a wrong answer. Earlier "
            "studies lumped these together, which conflates two opposing "
            "effects."
        ),
        "qa": [],
    },
    {
        "n": 15, "title": "Limitations",
        "screen": "Title - likely a bullet list of caveats.",
        "say": (
            "We're up-front about the limitations, and there are several.\n\n"
            "One. The success label is a textual heuristic. We can't run "
            "all six thousand pieces of code at scale, so a conversation "
            "may be labeled success even when the produced code is subtly "
            "wrong but the user said nothing. The kappa-0.834 audit gives "
            "us confidence the heuristic is sound, but it's not a "
            "substitute for execution.\n\n"
            "Two. The study is observational, not causal. Skilled "
            "developers may both write better prompts and choose easier "
            "tasks, and we can't fully rule that out. We control for "
            "language, source type, snapshot, and ChatGPT model variant, "
            "but unobserved confounders like developer experience are not "
            "in the data. So every claim we make is framed as an "
            "association after adjustment, not a causal effect.\n\n"
            "Three. DevGPT contains conversations the developer chose to "
            "share publicly, which introduces selection bias toward funny "
            "or dramatic threads.\n\n"
            "Four. The corpus is English-heavy and skews toward web "
            "languages like CSS, JavaScript, and TypeScript, which limits "
            "external validity to other language ecosystems.\n\n"
            "Five. The cross-vendor study uses a rubric judge that is "
            "itself an LLM. We mitigate by reporting effect sizes "
            "alongside p-values and using inter-vendor kappa as a sanity "
            "check, but a study with human raters would strengthen that "
            "piece.\n\n"
            "Six. The dataset window stops at May 2024, so newer ChatGPT "
            "versions are not represented."
        ),
        "qa": [],
    },
    {
        "n": 16, "title": "Conclusion",
        "screen": "Title - likely a recap.",
        "say": (
            "To summarize. We mined eleven structural and behavioral "
            "prompt features from 6,413 real developer ChatGPT "
            "conversations, fit a balanced logistic regression with "
            "five-fold stratified cross validation, and reached ROC-AUC "
            "0.799.\n\n"
            "Three findings carry the story. One, an explicit output "
            "format instruction raises the odds of a successful "
            "conversation by 75 percent. Two, role instructions hurt on "
            "their own but help by a factor of about five when combined "
            "with an output format. Three, the engineered prompt effect "
            "generalizes across three independent LLMs, with success "
            "rates rising 14 to 24 percentage points.\n\n"
            "Feature group ablation, pre-registered interactions, manual "
            "validation at kappa 0.834, per-snapshot stability, and the "
            "cross-vendor replication all point in the same direction. "
            "The structural prompt features carry roughly three times the "
            "predictive load of language, source type, model variant, and "
            "snapshot date combined.\n\n"
            "That's a converging body of evidence from four independent "
            "methods, on real data, with proper controls. The paper turns "
            "prompt-engineering folklore into a reproducible statistical "
            "claim."
        ),
        "qa": [],
    },
    {
        "n": 17, "title": "Future Work",
        "screen": "Title - likely four bullets.",
        "say": (
            "Four directions extend this study.\n\n"
            "First, a fully blinded two-reviewer human re-labeling pass "
            "on a larger sample of around 500 conversations. That would "
            "let us report human-against-human kappa separately from "
            "human-against-auto kappa, tightening the validation argument.\n\n"
            "Second, a causal study using natural experiments inside the "
            "DevGPT timeline - for example, the GPT-3.5 to GPT-4 model "
            "cutover. That would move the work from association to "
            "intervention.\n\n"
            "Third, an extension to non-English developer corpora and to "
            "languages outside the JavaScript and Python clusters. That "
            "would test external validity.\n\n"
            "Fourth, a fine-grained interaction screen with Bonferroni or "
            "Benjamini-Hochberg correction over all eleven choose two "
            "possible two-way terms - that's 55 pairs. Pre-registering "
            "only four interactions was a deliberate, conservative choice "
            "to avoid false positives, but it leaves potential compounding "
            "effects un-screened. A multiple-comparison-corrected "
            "exploratory pass would expose any we missed."
        ),
        "qa": [],
    },
    {
        "n": 18, "title": "References (page 1)",
        "screen": "References [1] through [8].",
        "say": (
            "These are the first eight references. Three I want to call "
            "out by name. Reference 6 is the DevGPT v10 release by Xiao "
            "and colleagues - that's the public corpus that made this "
            "work possible. Reference 1, Mohamed et al., is the closest "
            "descriptive precursor to our paper. And reference 8, Prompt "
            "Alchemy by Ye and colleagues, is one of the automated "
            "prompt-refinement systems we cite as related work.\n\n"
            "I'll skip the rest unless anyone has a question about a "
            "specific citation.\n\n"
            "[Pacing: 15-20 seconds. The audience won't read references; "
            "they want you to move on.]"
        ),
        "qa": [],
    },
    {
        "n": 19, "title": "References (page 2)",
        "screen": "References [9] through [15].",
        "say": (
            "Continuation of the reference list. The two I'll mention are "
            "reference 13 - Landis and Koch 1977, which is the kappa-"
            "interpretation rubric we use throughout - and references 14 "
            "and 15, which are the scikit-learn and statsmodels papers, "
            "since those are the libraries that powered our regression.\n\n"
            "That covers the bibliography.\n\n"
            "[Pacing: 10-15 seconds.]"
        ),
        "qa": [],
    },
    {
        "n": 20, "title": "Thank You",
        "screen": "Thank You!!!",
        "say": (
            "That brings us to the end. To recap one more time: prompt "
            "features predict success in real developer-ChatGPT "
            "conversations, reproducibly and across vendors, with an "
            "output-format instruction giving a 75 percent lift on its own "
            "and a role-times-format interaction giving roughly a "
            "five-times boost.\n\n"
            "The paper, the code, the dataset processing, and a live "
            "companion website are all in our GitHub repository. We'd be "
            "glad to share the URL after the talk.\n\n"
            "Thank you. We're happy to take questions."
        ),
        "qa": [],
    },
]

UNIVERSAL_QA = [
    (
        "Why didn't you use a transformer-based classifier instead?",
        "We did fit a 2,000-feature bag-of-words logistic as a stronger "
        "baseline. It hits AUC 0.782, so we're 1.7 points above it. A "
        "transformer would likely add another point or two, but we'd lose "
        "interpretability, which is the whole point - we want to be able "
        "to say which feature matters and by how much. The eleven "
        "engineered features pay for themselves: ablation shows they carry "
        "three times the predictive load of all controls combined."
    ),
    (
        "How do you know the auto-labeler isn't biased?",
        "Three pieces of evidence. (1) Cohen's kappa of 0.834 against "
        "manual labels - almost perfect on Landis and Koch. (2) The four "
        "disagreements are all conservative - the auto-labeler is stricter "
        "than the human, so any bias goes against finding effects. "
        "(3) The cross-model study uses a separate rubric and still finds "
        "the same engineered-prompt lift, so the result is not driven by "
        "auto-labeler quirks."
    ),
    (
        "What's the most surprising finding for you?",
        "The role-times-format interaction. Bare role hurts on its own, "
        "but the combination with format flips strongly positive. That's "
        "a textbook synergy and it's what reconciles the popular wisdom "
        "of 'you are an X expert' prompting with the data showing it "
        "doesn't work without an anchor."
    ),
    (
        "Can you share the dataset?",
        "DevGPT v10 is already public - Xiao et al., arXiv:2309.03914. "
        "Our processing pipeline, regex catalog, and feature parquet "
        "files are in our GitHub repo, MIT-licensed."
    ),
    (
        "What's the practical advice for someone in this room writing prompts tomorrow morning?",
        "Three things. Add an output format anchor - JSON, markdown, "
        "fenced code, anything concrete. Don't write a bare role prompt "
        "without that anchor. And don't be afraid to refine - "
        "back-and-forth is positive, as long as you're not in a "
        "correction loop fixing wrong answers."
    ),
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_label_paragraph(doc, label, body):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3B, 0x4D, 0x80)
    p.add_run("  ")
    body_run = p.add_run(body)
    body_run.font.size = Pt(11)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Speaker Notes - Prompt Pattern Mining in Vibe Coding", level=0
    )
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "Per-slide delivery script for the Graduate Seminar presentation. "
        "Total target time: 18-22 minutes for 19 content slides plus thanks. "
        "Each slide block lists what is on the slide, the speaker text to "
        "deliver, and the most likely audience question with a clean answer."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    doc.add_paragraph()

    for note in NOTES:
        doc.add_page_break()
        add_heading(
            doc, f"Slide {note['n']} - {note['title']}", level=1
        )
        add_label_paragraph(doc, "On screen:", note["screen"])
        doc.add_paragraph()
        say_h = doc.add_paragraph()
        say_run = say_h.add_run("Say")
        say_run.bold = True
        say_run.font.size = Pt(12)
        say_run.font.color.rgb = RGBColor(0x14, 0x55, 0x14)
        for chunk in note["say"].split("\n\n"):
            doc.add_paragraph(chunk)
        if note["qa"]:
            doc.add_paragraph()
            qa_h = doc.add_paragraph()
            qa_run = qa_h.add_run("If asked")
            qa_run.bold = True
            qa_run.font.size = Pt(12)
            qa_run.font.color.rgb = RGBColor(0x80, 0x3B, 0x14)
            for q, a in note["qa"]:
                pq = doc.add_paragraph()
                pqr = pq.add_run("Q: ")
                pqr.bold = True
                pq.add_run(q)
                pa = doc.add_paragraph()
                par = pa.add_run("A: ")
                par.bold = True
                pa.add_run(a)

    doc.add_page_break()
    add_heading(doc, "Universal Q&A backup", level=1)
    doc.add_paragraph(
        "These come up often. Memorize them so you can answer in one breath."
    )
    for q, a in UNIVERSAL_QA:
        pq = doc.add_paragraph()
        pqr = pq.add_run("Q: ")
        pqr.bold = True
        pq.add_run(q)
        pa = doc.add_paragraph()
        par = pa.add_run("A: ")
        par.bold = True
        pa.add_run(a)
        doc.add_paragraph()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
